import streamlit as st
import torch
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import docx
import PyPDF2
import re
import requests
from newspaper import Article
from pathlib import Path
import time

# Các tham số cố định

DEFAULT_MODEL_DIR = "outputs/bartpho-finetuned"

# File test mặc định trong mô trường
DEFAULT_TEST_DOCX = "/mnt/data/baocaohocmay.docx"

st.set_page_config(page_title="BartPho Summarizer — Clean UI", layout="wide")


# Hàm hỗ trợ đọc file pdf
def read_pdf(file_obj):
    """
    Đọc và trích xuất nội dung từ file PDF.

    Parameters
    ----------
    file_obj : UploadedFile
        File PDF được upload qua Streamlit.

    Returns
    -------
    str
        Nội dung text đã được ghép từ các trang PDF.
    """
    reader = PyPDF2.PdfReader(file_obj)
    text = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text.append(page_text)
    return "\n".join(text)

# Hàm hỗ trợ đọc file docx
def read_docx(path_or_file):
    """
    Đọc nội dung từ file DOCX.  
    Hỗ trợ cả đường dẫn file và file upload từ Streamlit.

    Parameters
    ----------
    path_or_file : str or UploadedFile
        Đường dẫn file hoặc đối tượng file upload.

    Returns
    -------
    str
        Nội dung văn bản thuần (text) từ file DOCX.
    """
    if hasattr(path_or_file, "read"):
        doc = docx.Document(path_or_file)
    else:
        doc = docx.Document(str(path_or_file))
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def clean_text(text: str) -> str:
    """
    Làm sạch văn bản: loại bỏ tab, khoảng trắng thừa, newline thừa.

    Parameters
    ----------
    text : str
        Văn bản cần làm sạch.

    Returns
    -------
    str
        Văn bản đã được chuẩn hóa.
    """
    if not text:
        return ""
    text = text.strip()
    text = re.sub(r"\t+", " ", text)
    text = re.sub(r" *\n+ *", "\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text


# Load model & tokenizer với cache để tránh load lại nhiều lần
@st.cache_resource
def load_model_and_tokenizer(model_dir: str):
    """
    Load tokenizer và mô hình tóm tắt từ thư mục local (đã fine-tune).

    Parameters
    ----------
    model_dir : str
        Thư mục chứa model và tokenizer.

    Returns
    -------
    tuple
        (tokenizer, model, device) đã được load và đưa lên GPU nếu có.
    """
    tokenizer = AutoTokenizer.from_pretrained(model_dir)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_dir)
    device = "cuda" if torch.cuda.is_available() else "cpu"
    model = model.to(device)
    return tokenizer, model, device


# Chia văn bản thành các đoạn theo token limit
def chunk_text_by_tokens(text: str, tokenizer, max_tokens: int = 800, overlap: int = 64):
    """
    Chia văn bản dài thành nhiều chunk dựa trên số lượng token.

    Parameters
    ----------
    text : str
        Nội dung văn bản gốc.
    tokenizer : AutoTokenizer
        Tokenizer để encode text thành token IDs.
    max_tokens : int
        Số token tối đa cho mỗi chunk.
    overlap : int
        Số token lặp lại giữa các chunk (giúp giữ ngữ cảnh).

    Returns
    -------
    list[str]
        Danh sách các chunk đã tách.
    """
    ids = tokenizer.encode(text)
    if len(ids) <= max_tokens:
        return [text]
    chunks = []
    start = 0
    while start < len(ids):
        end = start + max_tokens
        chunk_ids = ids[start:end]
        chunk_text = tokenizer.decode(chunk_ids, skip_special_tokens=True, clean_up_tokenization_spaces=True)
        chunk_text = clean_text(chunk_text)
        if chunk_text:
            chunks.append(chunk_text)
        start = max(0, end - overlap)
        if end >= len(ids):
            break
    return chunks


# Hàm tóm tắt từng chunk
def summarize_chunk(chunk: str, tokenizer, model, device, max_summary_tokens=128, num_beams=4):
    """
    Tóm tắt một chunk nội dung.

    Parameters
    ----------
    chunk : str
        Phần văn bản cần tóm tắt.
    tokenizer : AutoTokenizer
        Tokenizer đã load.
    model : AutoModelForSeq2SeqLM
        Mô hình đã load.
    device : str
        'cpu' hoặc 'cuda'.
    max_summary_tokens : int
        Độ dài tối đa của bản tóm tắt.
    num_beams : int
        Beam search width.

    Returns
    -------
    str
        Văn bản tóm tắt đã làm sạch.
    """
    inputs = tokenizer(chunk, return_tensors="pt", truncation=True, max_length=1024)
    input_ids = inputs["input_ids"].to(device)
    attention_mask = inputs.get("attention_mask", None)
    gen_kwargs = dict(max_length=max_summary_tokens, num_beams=num_beams, early_stopping=True)
    if attention_mask is not None:
        gen = model.generate(input_ids=input_ids, attention_mask=attention_mask.to(device), **gen_kwargs)
    else:
        gen = model.generate(input_ids=input_ids, **gen_kwargs)
    summary = tokenizer.decode(gen[0], skip_special_tokens=True)
    return clean_text(summary)


# Lấy nội dung từ url
def extract_text_from_url(url: str, timeout: int = 10):
    """
    Lấy nội dung text từ URL qua Newspaper3k, fallback sang HTML parsing thô.

    Parameters
    ----------
    url : str
        Địa chỉ trang web.
    timeout : int
        Timeout tải nội dung.

    Returns
    -------
    str
        Nội dung text đã làm sạch.

    Raises
    ------
    RuntimeError
        Nếu không thể tải nội dung từ URL.
    """
    try:
        article = Article(url)
        article.download()
        article.parse()
        txt = clean_text(article.text)
        if txt and len(txt) > 50:
            return txt
    except Exception:
        pass
    # fallback: bóc html thô
    try:
        r = requests.get(url, timeout=timeout)
        html = r.text
        text = re.sub(r"<script.*?>.*?</script>", "", html, flags=re.S)
        text = re.sub(r"<style.*?>.*?</style>", "", text, flags=re.S)
        text = re.sub(r"<[^>]+>", " ", text)
        return clean_text(text)
    except Exception as e:
        raise RuntimeError(f"Không thể tải nội dung từ URL: {e}")


# Layout chính của ứng dụng Streamlit
with st.sidebar:
    st.header("Cấu hình tóm tắt")
    model_dir = st.text_input("Model directory", value=DEFAULT_MODEL_DIR)
    max_chunk_tokens = st.number_input("Max chunk tokens", min_value=128, max_value=2048, value=800, step=64)
    chunk_overlap = st.number_input("Chunk overlap (tokens)", min_value=0, max_value=512, value=64, step=16)
    max_summary_tokens = st.number_input("Max summary tokens per chunk", min_value=32, max_value=512, value=128, step=16)
    num_beams = st.slider("Beams (num_beams)", 1, 8, 4)
    run_button = st.button("(Re)Load model")
    st.markdown("---")
    st.markdown("Ví dụ: bạn đã upload 1 file test: **/mnt/data/baocaohocmay.docx**")

# load model nếu lần đầu hoặc khi nhấn reload
if "model_loaded" not in st.session_state or run_button:
    try:
        with st.spinner("Đang load tokenizer & model (có thể mất vài giây)..."):
            tokenizer, model, device = load_model_and_tokenizer(model_dir)
        st.session_state.tokenizer = tokenizer
        st.session_state.model = model
        st.session_state.device = device
        st.session_state.model_loaded = True
        st.success("Model loaded ✔")
    except Exception as e:
        st.error(f"Lỗi khi load model: {e}")
        st.stop()
else:
    tokenizer = st.session_state.tokenizer
    model = st.session_state.model
    device = st.session_state.device

# Upload file hoặc nhập url
st.title("BartPho — Summarizer (refined)")
col1, col2 = st.columns([2, 1])

with col1:
    st.subheader("1) Nhập nội dung")
    uploaded = st.file_uploader("Upload PDF / DOCX (hoặc để trống dùng file mẫu)", type=["pdf", "docx"], accept_multiple_files=False)
    url_input = st.text_input("Hoặc nhập URL bài báo để tóm tắt")
    use_sample = st.checkbox("Sử dụng file mẫu (/mnt/data/baocaohocmay.docx)", value=False)

with col2:
    st.subheader("2) Tùy chọn nhanh")
    show_orig = st.checkbox("Hiện nội dung gốc (cắt ngắn)", value=True)
    show_chunks = st.checkbox("Hiện các chunk", value=False)
    show_progress = st.checkbox("Hiện progress bar", value=True)

# Thu thập nội dung từ nguồn tương ứng
text = ""
source_label = ""
if uploaded:
    source_label = uploaded.name
    if uploaded.type == "application/pdf":
        text = read_pdf(uploaded)
    else:
        text = read_docx(uploaded)
elif url_input:
    source_label = url_input
    try:
        with st.spinner("Đang tải nội dung từ URL ..."):
            text = extract_text_from_url(url_input)
    except Exception as e:
        st.error(str(e))
        st.stop()
elif use_sample:
    path = Path(DEFAULT_TEST_DOCX)
    if path.exists():
        source_label = str(path)
        text = read_docx(path)
    else:
        st.error(f"File mẫu không tồn tại: {path}")
        st.stop()
else:
    st.info("Vui lòng upload file hoặc nhập URL hoặc chọn file mẫu để bắt đầu.")

if not text:
    st.stop()

text = clean_text(text)

# Hiển thị nội dung nguồn
if show_orig:
    st.subheader(f"Nội dung nguồn — {source_label}")
    st.write(text[:3000] + (" ..." if len(text) > 3000 else ""))

# chunking
chunks = chunk_text_by_tokens(text, tokenizer, max_tokens=int(max_chunk_tokens), overlap=int(chunk_overlap))

st.write(f"🔹 Tổng tokens (ước lượng): {len(tokenizer.encode(text))} — Số chunk: {len(chunks)}")
if show_chunks:
    for i, c in enumerate(chunks, 1):
        st.markdown(f"**Chunk {i}** — ({len(tokenizer.encode(c))} tokens)")
        st.write(c[:1000] + (" ..." if len(c) > 1000 else ""))

# Summarize button
if st.button("Tóm tắt now — Summarize"):
    final = []
    progress = st.progress(0)
    total = len(chunks)
    for i, ch in enumerate(chunks, 1):
        if show_progress:
            st.write(f"⏳ Tóm tắt chunk {i}/{total} — tokens {len(tokenizer.encode(ch))}")
        try:
            s = summarize_chunk(ch, tokenizer, model, device, max_summary_tokens=int(max_summary_tokens), num_beams=int(num_beams))
        except Exception as e:
            s = f"⚠ Lỗi khi tóm tắt chunk {i}: {e}"
        final.append(s)
        if show_progress:
            progress.progress(int(i/total*100))
    joined = "\n\n".join(final)
    st.subheader("Kết quả tóm tắt (từng chunk ghép lại)")
    st.write(clean_text(joined))
    # Final short summary
    if len(tokenizer.encode(joined)) > 32:
        with st.expander("Tạo tóm tắt ngắn gọn từ bản tóm tắt" , expanded=False):
            try:
                short = summarize_chunk(joined, tokenizer, model, device, max_summary_tokens=200, num_beams=4)
                st.markdown("**Tóm tắt ngắn gọn:**")
                st.write(short)
            except Exception as e:
                st.write(f"Không thể tạo tóm tắt ngắn hơn: {e}")


# Footer
st.markdown("---")
st.caption("Ứng dụng được tối ưu để chạy với mô hình fine-tuned tại local. Cần đảm bảo model và tokenizer nằm trong thư mục cấu hình.")
